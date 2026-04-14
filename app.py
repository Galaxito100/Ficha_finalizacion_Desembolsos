import re
import os
import io
import zipfile
import tempfile
import streamlit as st
from pathlib import Path
from cryptography.fernet import Fernet
import json
import base64
import hashlib

# ── Configuración de página ────────────────────────────────────────────────────
st.set_page_config(
    page_title="CAF – Extractor FR",
    layout="wide",
    page_icon="🏛️",
    initial_sidebar_state="collapsed"
)

# ── CLASE PARA CIFRADO SEGURO (Se mantiene para proteger los datos) ────────────
class SecureData:
    """Cifra y descifra datos sensibles de la aplicación"""
    
    def __init__(self):
        try:
            clave_raw = st.secrets["security"]["encryption_key"]
            
            if isinstance(clave_raw, str):
                clave_raw = clave_raw.encode()
            
            if len(clave_raw) < 32:
                clave_raw = clave_raw.ljust(32, b'\0')
            elif len(clave_raw) > 32:
                clave_raw = hashlib.sha256(clave_raw).digest()
            
            clave_base64 = base64.urlsafe_b64encode(clave_raw)
            self.cipher = Fernet(clave_base64)
            self.disponible = True
            
        except Exception as e:
            st.error(f"⚠️ Error configurando cifrado: {e}")
            self.disponible = False
    
    def cifrar_datos(self, datos):
        if not self.disponible:
            return datos
        try:
            json_str = json.dumps(datos, ensure_ascii=False)
            datos_cifrados = self.cipher.encrypt(json_str.encode('utf-8'))
            return base64.b64encode(datos_cifrados).decode('utf-8')
        except Exception as e:
            st.error(f"Error cifrando datos: {e}")
            return None
    
    def descifrar_datos(self, datos_cifrados):
        if not self.disponible:
            return datos_cifrados
        if not datos_cifrados:
            return None
        try:
            datos_bytes = base64.b64decode(datos_cifrados)
            datos_descifrados = self.cipher.decrypt(datos_bytes)
            return json.loads(datos_descifrados.decode('utf-8'))
        except Exception as e:
            return None
    
    def cifrar_texto(self, texto):
        if not self.disponible or not texto:
            return texto
        try:
            datos_cifrados = self.cipher.encrypt(texto.encode('utf-8'))
            return base64.b64encode(datos_cifrados).decode('utf-8')
        except:
            return texto
    
    def descifrar_texto(self, texto_cifrado):
        if not self.disponible or not texto_cifrado:
            return texto_cifrado
        try:
            datos_bytes = base64.b64decode(texto_cifrado)
            return self.cipher.decrypt(datos_bytes).decode('utf-8')
        except:
            return texto_cifrado

# ── INICIALIZAR COMPONENTES ─────────────────────────────────────────────────
if 'secure' not in st.session_state:
    st.session_state['secure'] = SecureData()

secure = st.session_state['secure']

# ── ESTILO CAF ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;700;800&family=Open+Sans:wght@400;600&display=swap');

html, body, [class*="css"] {
    font-family: 'Open Sans', sans-serif;
    background-color: #f4f6f9;
}

.caf-header {
    background: linear-gradient(135deg, #004A8F 0%, #006BB6 100%);
    padding: 28px 36px;
    border-radius: 12px;
    margin-bottom: 28px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    box-shadow: 0 4px 20px rgba(0,74,143,0.25);
}
.caf-header-title {
    color: white;
    font-family: 'Montserrat', sans-serif;
    font-size: 26px;
    font-weight: 800;
    letter-spacing: 0.5px;
    margin: 0;
}
.caf-header-sub {
    color: rgba(255,255,255,0.80);
    font-size: 13px;
    font-weight: 600;
    letter-spacing: 1.5px;
    text-transform: uppercase;
    margin-top: 4px;
}
.section-header {
    background: linear-gradient(90deg, #004A8F 0%, #006BB6 100%);
    color: white;
    font-family: 'Montserrat', sans-serif;
    font-weight: 700;
    font-size: 14px;
    letter-spacing: 1.2px;
    text-transform: uppercase;
    padding: 10px 20px;
    border-radius: 6px;
    margin: 24px 0 0 0;
}
.caf-table {
    width: 100%;
    border-collapse: collapse;
    margin-top: 12px;
    margin-bottom: 8px;
    border-radius: 8px;
    overflow: hidden;
    box-shadow: 0 2px 12px rgba(0,0,0,0.08);
}
.caf-table tr:nth-child(even) { background-color: #EEF3F9; }
.caf-table tr:nth-child(odd)  { background-color: #ffffff; }
.caf-table td {
    padding: 11px 18px;
    font-size: 13.5px;
    border-bottom: 1px solid #dce6f0;
    vertical-align: top;
}
.caf-table td:first-child {
    background-color: #004A8F;
    color: white;
    font-family: 'Montserrat', sans-serif;
    font-weight: 600;
    font-size: 12.5px;
    width: 280px;
    letter-spacing: 0.3px;
}
.caf-table td:last-child { color: #1a2e45; }
div.stButton > button {
    background: linear-gradient(135deg, #004A8F, #006BB6);
    color: white;
    font-family: 'Montserrat', sans-serif;
    font-weight: 700;
    font-size: 14px;
    letter-spacing: 0.8px;
    border: none;
    padding: 12px 36px;
    border-radius: 8px;
    cursor: pointer;
    transition: all 0.2s ease;
    box-shadow: 0 4px 14px rgba(0,74,143,0.3);
    width: 100%;
}
div.stButton > button:hover {
    transform: translateY(-1px);
    box-shadow: 0 6px 18px rgba(0,74,143,0.4);
}
.caf-footer {
    text-align: center;
    color: #8fa3bd;
    font-size: 11.5px;
    margin-top: 40px;
    padding-top: 16px;
    border-top: 1px solid #dce6f0;
    letter-spacing: 0.5px;
}
#MainMenu, footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ── Header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="caf-header">
    <div>
        <div class="caf-header-sub">Gerencia Corporativa de Riesgos · Dirección de Riesgo Soberano</div>
        <div class="caf-header-title">Ficha de Finalización de Desembolsos</div>
    </div>
    <div style="color:white; font-family:Montserrat; font-size:32px; font-weight:900; letter-spacing:-1px;">CAF</div>
</div>
""", unsafe_allow_html=True)

# ── Upload + Botón ─────────────────────────────────────────────────────────────
col_up, col_btn = st.columns([3, 1])
with col_up:
    archivo = st.file_uploader("Sube tu documento (.pdf o .docx)", type=["pdf", "docx"])
with col_btn:
    st.write("")
    st.write("")
    procesar = st.button("⚙️ Procesar")

archivo_excel = st.file_uploader(
    "Sube la base de EED (.xlsx) para verificación de calidad",
    type=["xlsx"],
    help="Debe contener la hoja 'Consolidado (2019 - 2025)' con columna 'Codigo de documento'"
)

st.markdown('<div style="margin-top: 20px; font-size: 13.5px; color: #004A8F; font-weight: 600;">📚 Documentación de Respaldo (ZIP)</div>', unsafe_allow_html=True)
st.caption("Comprime todos tus PDFs de respaldo (EED, CCI, etc.) en un solo archivo .zip para procesamiento rápido.")
zip_pdfs = st.file_uploader(
    "Subir carpeta comprimida (.zip)",
    type=["zip"],
    help="Sube un archivo .zip que contenga todos los PDFs para contrastar"
)

# ── Funciones de extracción ────────────────────────────────────────────────────
def extraer_texto_pdf(ruta):
    import pdfplumber
    texto_completo = []
    with pdfplumber.open(ruta) as pdf:
        for pagina in pdf.pages:
            for tabla in pagina.extract_tables():
                for fila in tabla:
                    texto_completo.append(" | ".join([c or "" for c in fila]))
            texto_plano = pagina.extract_text()
            if texto_plano:
                texto_completo.append(texto_plano)
    return "\n".join(texto_completo)

def extraer_texto_docx(ruta):
    from docx import Document
    doc = Document(ruta)
    lineas = []
    for parrafo in doc.paragraphs:
        if parrafo.text.strip():
            lineas.append(parrafo.text.strip())
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            celdas = list(dict.fromkeys(celdas))
            lineas.append(" | ".join(celdas))
    return "\n".join(lineas)

SEP = r"[\s|]+"

def buscar_campo(texto, patron, grupo=1):
    match = re.search(patron, texto, re.IGNORECASE | re.MULTILINE)
    if not match:
        return "No encontrado"
    valor = match.group(grupo).strip().strip("|").strip()
    return valor if valor else "No encontrado"

def normalizar(t):
    import unicodedata
    return "".join(c for c in unicodedata.normalize("NFD", t.lower()) if unicodedata.category(c) != "Mn").strip()

def celda_a_html(celda):
    from docx.oxml.ns import qn
    import re as _re
    partes = []
    for parrafo in celda.paragraphs:
        texto_parrafo = ""
        elems = list(parrafo._p)
        skip_until_end = False
        url_campo = ""
        texto_campo = ""
        in_separate = False
        for elem in elems:
            tag = elem.tag.split("}")[-1]
            if tag == "hyperlink":
                rId = elem.get(qn("r:id"))
                url = ""
                if rId and rId in celda.part.rels:
                    url = celda.part.rels[rId].target_ref
                texto_link = "".join(
                    t.text for r in elem.findall(qn("w:r"))
                    for t in r.findall(qn("w:t")) if t.text
                )
                if url and texto_link:
                    texto_parrafo += f'<a href="{url}" target="_blank" style="color:#006BB6">{texto_link}</a>'
                else:
                    texto_parrafo += texto_link
            elif tag == "r":
                fld = elem.find(qn("w:fldChar"))
                instr = elem.find(qn("w:instrText"))
                t_elem = elem.find(qn("w:t"))
                if fld is not None:
                    ft = fld.get(qn("w:fldCharType"), "")
                    if ft == "begin":
                        url_campo = ""; texto_campo = ""; in_separate = False; skip_until_end = True
                    elif ft == "separate":
                        in_separate = True
                    elif ft == "end":
                        if url_campo and texto_campo:
                            texto_parrafo += f'<a href="{url_campo}" target="_blank" style="color:#006BB6">{texto_campo}</a>'
                        elif texto_campo:
                            texto_parrafo += texto_campo
                        skip_until_end = False; in_separate = False
                elif instr is not None and instr.text:
                    m = _re.search(r'HYPERLINK\s+"([^"]+)"', instr.text)
                    if m: url_campo = m.group(1)
                elif t_elem is not None and t_elem.text:
                    if in_separate: texto_campo += t_elem.text
                    elif not skip_until_end: texto_parrafo += t_elem.text
        if texto_parrafo.strip():
            partes.append(texto_parrafo)
    return "<br>".join(partes) if partes else ""

def extraer_presentacion_informes(ruta, extension):
    resultados = {"Última auditoría": "No encontrado", "Final": "No encontrado", "Pendientes": "No encontrado"}
    mapa = [
        ("ultima auditoria", "Última auditoría"), ("auditoria financiera", "Última auditoría"),
        ("auditor", "Última auditoría"), ("informe final del prestamo", "Final"),
        ("informe final", "Final"), ("final", "Final"), ("condiciones pendientes", "Pendientes"),
        ("condicion", "Pendientes"), ("pendientes", "Pendientes"), ("pendiente", "Pendientes"),
    ]
    try:
        if extension == ".docx":
            from docx import Document
            from docx.table import Table
            doc = Document(ruta)
            for elem in doc.element.body:
                if elem.tag.split("}")[-1] != "tbl": continue
                tabla = Table(elem, doc)
                texto_tabla = normalizar(" ".join(c.text for fila in tabla.rows for c in fila.cells))
                if "presentacion de informes" not in texto_tabla: continue
                
                fila_inicio = 0
                for j, fila in enumerate(tabla.rows):
                    texto_fila = normalizar(" ".join(c.text for c in fila.cells))
                    if "cumplimiento contractual" in texto_fila:
                        fila_inicio = j
                        break
                
                for j, fila in enumerate(tabla.rows):
                    if j < fila_inicio: continue
                    raw = [c.text.strip() for c in fila.cells]
                    seen = set(); celdas = []
                    for c in raw:
                        if c and c not in seen: seen.add(c); celdas.append(c)
                    if len(celdas) < 2: continue
                    etiqueta_norm = normalizar(celdas[-2])
                    valor_texto = celdas[-1].strip()
                    if not valor_texto or valor_texto in ("-", "—"): continue
                    
                    seen_ids = set(); celdas_unicas = []
                    for c in fila.cells:
                        cid = id(c._tc)
                        if cid not in seen_ids: seen_ids.add(cid); celdas_unicas.append(c)
                    
                    celda_contenido = celdas_unicas[-1] if len(celdas_unicas) >= 2 else None
                    valor_html = celda_a_html(celda_contenido) if celda_contenido else ""
                    valor = valor_html.strip() or valor_texto
                    for clave, nombre in mapa:
                        if clave in etiqueta_norm and resultados[nombre] == "No encontrado":
                            resultados[nombre] = valor
                            break
        else:
            import pdfplumber
            with pdfplumber.open(ruta) as pdf:
                texto_completo = "\n".join(p.extract_text() or "" for p in pdf.pages)
            m_sec = re.search(r"cumplimiento\s+contractual", texto_completo, re.IGNORECASE)
            texto = texto_completo[m_sec.start():] if m_sec else texto_completo
            for clave, nombre in mapa:
                m = re.search(rf"{re.escape(clave)}[^\n]*\n([^\n]+)", texto, re.IGNORECASE)
                if m and resultados[nombre] == "No encontrado":
                    resultados[nombre] = m.group(1).strip()
    except Exception as e:
        for k in resultados: resultados[k] = f"Error: {e}"
    return resultados

def extraer_objetivo_general(ruta, extension):
    try:
        if extension == ".docx":
            from docx import Document
            doc = Document(ruta)
            for tabla in doc.tables:
                for fila in tabla.rows:
                    celdas_raw = [c.text.strip() for c in fila.cells]
                    celdas = list(dict.fromkeys(celdas_raw))
                    if len(celdas) < 2: continue
                    if "objetivo general" not in celdas[0].lower(): continue
                    celda_valor = fila.cells[-1].text
                    m = re.search(r"Objetivo General[:\s]*\n+([\s\S]+?)(?:\n\s*\n|\nObjetivos Espec|$)", celda_valor, re.IGNORECASE)
                    if m: return " ".join(m.group(1).split())
                    if celda_valor.strip() and celda_valor.strip() not in ("-", "—"): return celda_valor.strip()
        else:
            import pdfplumber
            with pdfplumber.open(ruta) as pdf:
                texto = "\n".join(p.extract_text() or "" for p in pdf.pages)
            m = re.search(r"Objetivo General[:\s]*\n+([\s\S]+?)(?:\n\s*\n|Objetivos Espec)", texto, re.IGNORECASE)
            if m: return " ".join(m.group(1).split())
    except Exception: pass
    return "No encontrado"

def celda_dispensas_a_html(celda):
    from docx.oxml.ns import qn
    html = ""
    for elem in celda._tc:
        tag = elem.tag.split("}")[-1]
        if tag == "p":
            texto = "".join(n.text for n in elem.iter(qn("w:t")) if n.text).strip()
            if texto: html += f'<p style="margin:6px 0">{texto}</p>'
        elif tag == "tbl":
            from docx.table import Table
            tabla_interna = Table(elem, celda._tc)
            html += '<table style="border-collapse:collapse;width:100%;margin:10px 0">'
            for i, fila in enumerate(tabla_interna.rows):
                celdas = [c.text.strip() for c in fila.cells]
                celdas = list(dict.fromkeys(celdas))
                bg = "#EEF3F9" if i % 2 == 0 else "#ffffff"
                es_header = i == 0
                html += "<tr>"
                for c in celdas:
                    if es_header: html += f'<th style="background:#004A8F;color:white;padding:6px 10px;font-size:12px;text-align:left">{c}</th>'
                    else: html += f'<td style="background:{bg};padding:6px 10px;font-size:12px;border-bottom:1px solid #dce6f0">{c}</td>'
                html += "</tr>"
            html += "</table>"
    return html if html.strip() else None

def extraer_dispensas(ruta, extension):
    try:
        if extension == ".docx":
            from docx import Document
            doc = Document(ruta)
            for tabla in doc.tables:
                for fila in tabla.rows:
                    seen = set(); unicas = []
                    for c in fila.cells:
                        cid = id(c._tc)
                        if cid not in seen: seen.add(cid); unicas.append(c)
                    if len(unicas) < 2: continue
                    etiqueta = unicas[0].text.strip().lower()
                    if "dispensa" not in etiqueta and "enmienda" not in etiqueta: continue
                    celda = unicas[-1]
                    valor_texto = celda.text.strip()
                    if not valor_texto or valor_texto in ("-", "—"): continue
                    html = celda_dispensas_a_html(celda)
                    return html if html else valor_texto
        else:
            import pdfplumber
            with pdfplumber.open(ruta) as pdf:
                texto = "\n".join(p.extract_text() or "" for p in pdf.pages)
            m = re.search(r"Dispensas y enmiendas[\s|]+([\s\S]+?)(?:\n[A-Z][^\n]{0,40}\n)", texto, re.IGNORECASE)
            if m:
                parrafos = [p.strip() for p in m.group(1).splitlines() if p.strip()]
                return "<br>".join(parrafos)
    except Exception: pass
    return "No encontrado"

def extraer_codigos_de_zip(file_zip):
    """Extrae códigos EED/CCI/GOI/STCI de un archivo ZIP."""
    import pypdf
    codigos_encontrados = {}
    try:
        with zipfile.ZipFile(file_zip, 'r') as zf:
            pdf_files = [f for f in zf.namelist() if f.lower().endswith('.pdf')]
            for nombre_archivo in pdf_files:
                try:
                    with zf.open(nombre_archivo) as f:
                        reader = pypdf.PdfReader(f)
                        texto_completo = ""
                        for page in reader.pages:
                            texto_completo += page.extract_text() + "\n"
                    codigos = set(re.findall(r'((?:EED|CCI|GOI|STCI)[-\s][\d]+(?:/[\d]+)?)', texto_completo, re.I))
                    codigos_encontrados[nombre_archivo] = sorted(list(codigos))
                except Exception as e:
                    codigos_encontrados[nombre_archivo] = f"Error lectura: {str(e)}"
    except Exception as e:
        st.error(f"Error al descomprimir o leer el ZIP: {e}")
    return codigos_encontrados

def tabla_html(filas):
    rows = ""
    for label, valor in filas:
        rows += f"<tr><td>{label}</td><td>{valor or '—'}</td></tr>"
    return f'<table class="caf-table">{rows}</table>'

# ── Limpiar si no hay archivo ─────────────────────────────────────────────────
if archivo is None:
    for k in ["resultados_encrypted", "informes_encrypted", "objetivo_encrypted", 
              "dispensas_encrypted", "vista", "codigos_pdfs_encrypted", "total_dispensas", "total_pdfs"]:
        st.session_state.pop(k, None)

# ── Procesamiento ──────────────────────────────────────────────────────────────
if procesar:
    if archivo is None:
        st.warning("⚠️ Por favor sube un archivo antes de procesar.")
    else:
        extension = Path(archivo.name).suffix.lower()

        with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as tmp:
            tmp.write(archivo.read())
            ruta_tmp = tmp.name

        with st.spinner("Procesando documento..."):
            texto     = extraer_texto_pdf(ruta_tmp) if extension == ".pdf" else extraer_texto_docx(ruta_tmp)
            informes  = extraer_presentacion_informes(ruta_tmp, extension)
            objetivo  = extraer_objetivo_general(ruta_tmp, extension)
            dispensas = extraer_dispensas(ruta_tmp, extension)
            
            _m = re.search(r">Total</td>\s*<td[^>]*>\s*(\d+)", dispensas, re.IGNORECASE)
            total_dispensas = int(_m.group(1)) if _m else None
            
            codigos_pdfs = {}
            total_pdfs = 0
            if zip_pdfs:
                with st.spinner("Extrayendo códigos del ZIP..."):
                    codigos_pdfs = extraer_codigos_de_zip(zip_pdfs)
                    total_pdfs = len([c for c in codigos_pdfs.values() if isinstance(c, list)])

        # Eliminación de archivos temporales
        os.unlink(ruta_tmp)

        resultados_dict = {
            "N° de Operación (CFA)": buscar_campo(texto, r"CFA\s*[–\-]\s*([\d]+(?:/[\d]+)*)"),
            "Nombre de la Operación": buscar_campo(texto, r"Nombre de la operaci[oó]n" + SEP + r"([^|\n]+)"),
            "Prestatario": buscar_campo(texto, r"Prestatario" + SEP + r"([^|\n]+)"),
            "País": buscar_campo(texto, r"Pa[ií]s" + SEP + r"([^|\n]+)"),
            "Garante": buscar_campo(texto, r"Garante" + SEP + r"([^|\n]+)"),
            "Monto préstamo CAF (Aprobado)": buscar_campo(texto, r"Monto del pr[eé]stamo[\s\S]*?aprobado CAF" + SEP + r"((?:Contractual[^|\n]+|(?:US\$|USD)\s*[\d.,]+[^|\n]*))"),
            "Monto préstamo CAF (Desembolsado)": (lambda: (
                lambda c1, c2: c1 if c1 != "No encontrado" else c2
            )(
                buscar_campo(texto, r"Desembolsado:\s*((?:US\$|USD)\s*[\d.,]+[^\n]*)"),
                buscar_campo(texto, r"Monto del pr[eé]stamo[^\n]*aprobado CAF[^\n]*\nMonto del pr[eé]stamo[^\n]*aprobado CAF[\s|]+([^\n]+)")
            ))(),
        }

        # Cifrado de datos
        st.session_state["resultados_encrypted"] = secure.cifrar_datos(resultados_dict)
        st.session_state["informes_encrypted"] = secure.cifrar_datos(informes)
        st.session_state["objetivo_encrypted"] = secure.cifrar_texto(objetivo)
        st.session_state["dispensas_encrypted"] = secure.cifrar_texto(dispensas)
        st.session_state["total_dispensas"] = total_dispensas
        st.session_state["total_pdfs"] = total_pdfs
        st.session_state["codigos_pdfs_encrypted"] = secure.cifrar_datos(codigos_pdfs)
        st.session_state["vista"] = "informe"

# ── Renderizado ────────────────────────────────────────────────────────────────
if "resultados_encrypted" in st.session_state:
    resultados = secure.descifrar_datos(st.session_state["resultados_encrypted"])
    informes = secure.descifrar_datos(st.session_state["informes_encrypted"])
    objetivo = secure.descifrar_texto(st.session_state["objetivo_encrypted"])
    dispensas = secure.descifrar_texto(st.session_state["dispensas_encrypted"])
    codigos_pdfs = secure.descifrar_datos(st.session_state.get("codigos_pdfs_encrypted", "{}"))
    total_pdfs = st.session_state.get("total_pdfs", 0)
else:
    resultados = {}
    informes = {}
    objetivo = ""
    dispensas = ""
    codigos_pdfs = {}
    total_pdfs = 0

if resultados:
    total_dispensas = st.session_state.get("total_dispensas")

    # Botones de vista
    st.write("")
    col_v1, col_v2 = st.columns(2)
    with col_v1:
        if st.button("📋  Informe de Apoyo", use_container_width=True):
            st.session_state["vista"] = "informe"
    with col_v2:
        if st.button("🔍  Verificación de Calidad de la Ficha", use_container_width=True):
            st.session_state["vista"] = "calidad"

    vista = st.session_state.get("vista", "informe")

    if vista == "informe":
        st.markdown('<div class="section-header">📋 &nbsp;Informe de la Operación</div>', unsafe_allow_html=True)
        st.markdown(tabla_html(list(resultados.items())), unsafe_allow_html=True)

        st.markdown('<div class="section-header">📄 &nbsp;Presentación de Informes</div>', unsafe_allow_html=True)
        st.markdown(tabla_html([
            ("Última auditoría", informes["Última auditoría"]),
            ("Final", informes["Final"]),
            ("Pendientes", informes["Pendientes"]),
        ]), unsafe_allow_html=True)

        st.markdown('<div class="section-header">📝 &nbsp;Descripción de la Operación</div>', unsafe_allow_html=True)
        st.markdown(tabla_html([
            ("Objetivo General", objetivo),
        ]), unsafe_allow_html=True)

    elif vista == "calidad":
        dispensas_html = re.sub(
            r'((EED|CCI|GOI|STCI)[-\s][\d]+(?:/[\d]+)?)',
            r'<span style="color:#006BB6;font-weight:700">\1</span>',
            dispensas
        )
        st.markdown('<div class="section-header">⚖️ &nbsp;Dispensas y Enmiendas</div>', unsafe_allow_html=True)
        st.markdown(f'<div style="background:white;padding:18px 24px;border-radius:8px;margin-top:12px;box-shadow:0 2px 12px rgba(0,0,0,0.08);font-size:13.5px;color:#1a2e45;line-height:1.8;">{dispensas_html}</div>', unsafe_allow_html=True)

        if archivo_excel is not None:
            st.markdown('<div class="section-header">📊 &nbsp;Verificación contra Base EED</div>', unsafe_allow_html=True)
            try:
                import pandas as pd
                df = pd.read_excel(archivo_excel, sheet_name="Consolidado (2019 - 2025)")
                col_codigo = next((c for c in df.columns if "codigo" in c.lower() and "doc" in c.lower()), None)
                col_operacion = next((c for c in df.columns if "operaci" in c.lower() and "n" in c.lower()), None)
                
                if col_codigo is None: st.warning("⚠️ No se encontró la columna 'Codigo de documento' en el Excel.")
                elif col_operacion is None: st.warning("⚠️ No se encontró la columna 'Número de la operación' en el Excel.")
                else:
                    cfa_raw = resultados.get("N° de Operación (CFA)", "")
                    numeros_cfa = [n.strip() for n in cfa_raw.split("/") if n.strip()]

                    col_op_upper = df[col_operacion].astype(str).str.upper()
                    if len(numeros_cfa) == 1:
                        mask = col_op_upper.str.contains(f"CFA0*{numeros_cfa[0]}\\b", na=False, regex=True)
                    else:
                        mask = col_op_upper.apply(
                            lambda v: all(str(n) in str(v) for n in numeros_cfa)
                        )
                    
                    df_filtrado = df[mask]
                    total_excel = df_filtrado[col_codigo].dropna().nunique() if not df_filtrado.empty else 0
                    
                    codigos_ficha = sorted(set(re.findall(r'((?:EED|CCI|GOI|STCI)[-\s][\d]+(?:/[\d]+)?)', dispensas, re.IGNORECASE)))
                    
                    todos_codigos_zip = set()
                    for cods in codigos_pdfs.values():
                        if isinstance(cods, list): todos_codigos_zip.update(cods)
                    todos_codigos_zip = sorted(list(todos_codigos_zip))
                    
                    codigos_excel = sorted(df_filtrado[col_codigo].dropna().unique().tolist()) if not df_filtrado.empty else []
                    
                    col_a, col_b, col_c = st.columns(3)
                    with col_a: st.metric("Total en Ficha", total_dispensas if total_dispensas is not None else "No encontrado")
                    with col_b: st.metric("Total en Excel", total_excel)
                    with col_c: st.metric("Archivos en ZIP", total_pdfs)
                    
                    if total_dispensas is not None:
                        if total_dispensas == total_excel:
                            st.markdown('<div style="background:#E8F5E9;border-left:5px solid #43A047;border-radius:6px;padding:14px 18px;margin-top:8px;font-size:14px;color:#1b5e20;font-weight:700">✅ Coinciden</div>', unsafe_allow_html=True)
                        else:
                            st.markdown('<div style="background:#FFF8E1;border-left:5px solid #F5A623;border-radius:6px;padding:14px 18px;margin-top:8px;font-size:13px;color:#5a3e00;font-weight:700">⚠️ La cantidad de EED en la Ficha y la Base de Datos no coincide</div>', unsafe_allow_html=True)

                    st.write("")
                    col_t1, col_t2 = st.columns(2)

                    with col_t1:
                        st.markdown('**📄 Códigos en Ficha**')
                        if codigos_ficha:
                            html = '<table style="width:100%;border-collapse:collapse;font-size:13px"><tr><th style="background:#004A8F;color:white;padding:7px 12px;text-align:left">Código</th></tr>'
                            for i, cod in enumerate(codigos_ficha):
                                bg = "#EEF3F9" if i % 2 == 0 else "#ffffff"
                                html += f'<tr><td style="background:{bg};padding:7px 12px;border-bottom:1px solid #dce6f0;color:#006BB6;font-weight:700">{cod}</td></tr>'
                            html += '</table>'
                            st.markdown(html, unsafe_allow_html=True)
                        else: st.info("No se encontraron códigos.")

                    with col_t2:
                        st.markdown('**📊 Códigos en Excel**')
                        if codigos_excel:
                            html = '<table style="width:100%;border-collapse:collapse;font-size:13px"><tr><th style="background:#004A8F;color:white;padding:7px 12px;text-align:left">Código</th></tr>'
                            for i, cod in enumerate(codigos_excel):
                                bg = "#EEF3F9" if i % 2 == 0 else "#ffffff"
                                html += f'<tr><td style="background:{bg};padding:7px 12px;border-bottom:1px solid #dce6f0;color:#006BB6;font-weight:700">{cod}</td></tr>'
                            html += '</table>'
                            st.markdown(html, unsafe_allow_html=True)
                        else: st.info("No se encontraron registros.")

                    if codigos_pdfs:
                        st.write("")
                        st.markdown('<div class="section-header">📑 &nbsp;Códigos en ZIP (PDFs)</div>', unsafe_allow_html=True)
                        if todos_codigos_zip:
                            col_z1, col_z2 = st.columns([1, 2])
                            with col_z1:
                                st.markdown(f'<div style="background:#EEF3F9;padding:15px;border-radius:8px;text-align:center"><div style="font-size:24px;font-weight:800;color:#004A8F">{len(todos_codigos_zip)}</div><div style="font-size:12px;text-transform:uppercase">Total únicos</div></div>', unsafe_allow_html=True)
                            with col_z2:
                                st.markdown('**Por archivo:**')
                                for nombre, cods in codigos_pdfs.items():
                                    if isinstance(cods, list) and cods:
                                        display_name = nombre.split('/')[-1]
                                        st.markdown(f'<div style="display:inline-block;background:#fff;padding:6px 10px;margin:2px;border-radius:4px;border:1px solid #dce6f0;font-size:11px">📄 <b>{display_name}</b> ({len(cods)} found)</div>', unsafe_allow_html=True)
                        else:
                            st.info("No se encontraron códigos (EED/CCI/GOI/STCI) en el ZIP.")

            except Exception as e:
                st.error(f"Error al leer el Excel: {e}")
        else:
            st.info("📂 Sube la base de EED (.xlsx) para comparar el total de dispensas.")

    st.markdown('<div class="caf-footer">CAF – Banco de Desarrollo de América Latina y el Caribe &nbsp;·&nbsp; Gerencia Corporativa de Riesgos - 🔐 Datos Cifrados</div>', unsafe_allow_html=True)

